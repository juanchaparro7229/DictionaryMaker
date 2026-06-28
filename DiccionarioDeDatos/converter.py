import pandas as pd
from docx import Document
from docx.shared import Inches
import os
import glob

def create_data_dictionary():
    # Configuración de carpetas
    input_folder = 'input'
    output_folder = 'output'
    
    # Buscar archivos Excel en la carpeta de entrada
    excel_files = glob.glob(os.path.join(input_folder, '*.xlsx'))
    
    # Filtrar archivos temporales de Excel (empiezan con ~$)
    excel_files = [f for f in excel_files if not os.path.basename(f).startswith('~$')]
    
    if not excel_files:
        print(f"No se encontraron archivos .xlsx en la carpeta '{input_folder}'")
        return

    for file_path in excel_files:
        file_name = os.path.basename(file_path)
        print(f"Procesando: {file_name}")
        
        # Leer Excel
        df = pd.read_excel(file_path)
        
        # Crear documento Word
        doc = Document()
        doc.add_heading('Diccionario de Datos', 0)
        
        # Agrupar por Tabla
        tables = df['Tabla'].unique()
        
        for table_name in tables:
            doc.add_heading(f"Tabla: {table_name}", level=1)
            
            # Filtrar columnas de esta tabla
            table_data = df[df['Tabla'] == table_name]
            
            # Definir columnas para el Word (sin 'Tabla' porque ya es el encabezado)
            display_columns = [
                'Columna', 'TipoDato', 'LongitudMaxima', 'PermiteNulos', 
                'EsPK', 'EsFK', 'TablaRelacionada', 'Descripción'
            ]
            
            # Asegurar que 'Descripción' exista en el DataFrame (vacía)
            if 'Descripción' not in table_data.columns:
                table_data = table_data.assign(Descripción='')
            
            # Crear tabla en Word
            # +1 para el encabezado
            word_table = doc.add_table(rows=1, cols=len(display_columns))
            word_table.style = 'Table Grid'
            
            # Encabezados
            hdr_cells = word_table.rows[0].cells
            for i, col_name in enumerate(display_columns):
                hdr_cells[i].text = col_name
            
            # Datos
            for _, row in table_data.iterrows():
                row_cells = word_table.add_row().cells
                for i, col_name in enumerate(display_columns):
                    # Manejar valores nulos para evitar errores al escribir en Word
                    val = row.get(col_name, '')
                    if pd.isna(val):
                        val = ''
                    row_cells[i].text = str(val)
            
            doc.add_paragraph() # Espacio entre tablas

        # Guardar el documento
        output_file_name = os.path.splitext(file_name)[0] + '.docx'
        output_path = os.path.join(output_folder, output_file_name)
        doc.save(output_path)
        print(f"Archivo generado: {output_path}")

if __name__ == "__main__":
    create_data_dictionary()
